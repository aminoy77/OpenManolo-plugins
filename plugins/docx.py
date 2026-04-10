PLUGIN_NAME = "docx"
PLUGIN_DESCRIPTION = "Creates a real Microsoft Word .docx file with title and content sections"
PLUGIN_SCHEMA = {
    "type": "function",
    "function": {
        "name": "docx",
        "description": PLUGIN_DESCRIPTION,
        "parameters": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Full path where to save the .docx file"
                },
                "title": {
                    "type": "string",
                    "description": "Document title"
                },
                "sections": {
                    "type": "array",
                    "description": "List of sections, each with 'heading' and 'content'",
                    "items": {
                        "type": "object",
                        "properties": {
                            "heading": {"type": "string"},
                            "content": {"type": "string"}
                        }
                    }
                }
            },
            "required": ["path", "title", "sections"]
        }
    }
}


def run(path: str, title: str, sections: list) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Título
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Secciones
    for section in sections:
        heading = section.get("heading", "")
        content = section.get("content", "")

        if heading:
            doc.add_heading(heading, level=1)

        if content:
            for line in content.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

        doc.add_paragraph()

    doc.save(path)
    return f"Word document created: {path}"
